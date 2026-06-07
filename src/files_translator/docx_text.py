"""Extract paragraph text from DOCX files."""

from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import BinaryIO, Union

from docx import Document


@dataclass(frozen=True)
class ParagraphText:
    index: int
    text: str


def extract_paragraphs(source: Union[str, Path, BinaryIO, bytes]) -> list[ParagraphText]:
    """Return non-empty paragraphs with stable indices."""
    if isinstance(source, bytes):
        doc = Document(BytesIO(source))
    elif isinstance(source, (str, Path)):
        doc = Document(str(source))
    else:
        doc = Document(source)

    paragraphs: list[ParagraphText] = []
    for index, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            paragraphs.append(ParagraphText(index=index, text=text))
    return paragraphs
